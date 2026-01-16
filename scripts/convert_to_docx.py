import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Handle Tables
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Skip separator lines |---|---|
            if re.match(r'^\|[\s:-|]+\|$', line):
                continue
                
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table:
                # Process collected table data
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            table.cell(i, j).text = cell_text
                in_table = False
                table_data = []

        # Handle Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=2)
        
        # Handle Bullet Points
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle Empty Lines
        elif not line:
            continue
            
        # Handle Normal Paragraphs
        else:
            # Simple bold/italic parsing (optional, but requested for professional look)
            p = doc.add_paragraph()
            # Replace markdown bold **text** with rich text
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    md_file = r'c:\Users\LENOVO\.gemini\antigravity\brain\4aea0ce2-b9f9-4bf8-93e7-d84b724f9667\rapport_integration_fonctionnelle.md'
    docx_file = r'c:\Users\LENOVO\.gemini\antigravity\brain\4aea0ce2-b9f9-4bf8-93e7-d84b724f9667\Rapport_Integration_OptiChain.docx'
    md_to_docx(md_file, docx_file)
